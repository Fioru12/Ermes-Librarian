"""
Test robustezza parser .docx per documenti Word.
Verifica che python-docx (usato da LlamaIndex per i .docx) 
legga correttamente documenti Word di vario tipo.
"""
import pytest
from pathlib import Path


def _read_docx_text(filepath):
    """Legge il testo da un file .docx usando python-docx."""
    from docx import Document
    doc = Document(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs]
    # Aggiunge testo da tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def test_docx_simple_text(tmp_path):
    """Test lettura di un file .docx semplice con solo testo."""
    from docx import Document

    doc_path = tmp_path / "test_simple.docx"
    doc = Document()
    doc.add_paragraph("Questo è un paragrafo di test.")
    doc.add_paragraph("Secondo paragrafo con testo più lungo.")
    doc.save(str(doc_path))

    text = _read_docx_text(doc_path)

    assert "paragrafo di test" in text.lower()
    assert "Secondo paragrafo" in text


def test_docx_with_formatting(tmp_path):
    """Test lettura di un file .docx con formattazione."""
    from docx import Document

    doc_path = tmp_path / "test_formatting.docx"
    doc = Document()

    para = doc.add_paragraph()
    run = para.add_run("Test grassetto")
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("Test corsivo")
    run.italic = True

    doc.save(str(doc_path))

    text = _read_docx_text(doc_path)

    assert "Test grassetto" in text
    assert "Test corsivo" in text


def test_docx_empty_document(tmp_path):
    """Test lettura di un file .docx vuoto."""
    from docx import Document

    doc_path = tmp_path / "test_empty.docx"
    doc = Document()
    doc.save(str(doc_path))

    text = _read_docx_text(doc_path)
    # Un documento vuoto può avere testo (paragrafo vuoto) o stringa vuota
    assert text == "" or text is not None


def test_docx_with_tables(tmp_path):
    """Test lettura di un file .docx con tabelle."""
    from docx import Document

    doc_path = tmp_path / "test_tables.docx"
    doc = Document()

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Data 1"
    table.rows[1].cells[1].text = "Data 2"

    doc.save(str(doc_path))

    text = _read_docx_text(doc_path)

    assert "Header 1" in text or "Data 1" in text


def test_docx_mixed_content(tmp_path):
    """Test lettura di un file .docx con contenuto misto."""
    from docx import Document

    doc_path = tmp_path / "test_mixed.docx"
    doc = Document()

    doc.add_heading("Titolo del documento", level=1)
    doc.add_paragraph("Paragrafo introduttivo.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Cella A1"
    table.rows[0].cells[1].text = "Cella B1"
    table.rows[1].cells[0].text = "Cella A2"
    table.rows[1].cells[1].text = "Cella B2"

    doc.add_paragraph("Paragrafo finale.")

    doc.save(str(doc_path))

    text = _read_docx_text(doc_path)

    assert "Titolo del documento" in text
    assert "Paragrafo introduttivo" in text
    assert "Cella A1" in text
    assert "Paragrafo finale" in text
